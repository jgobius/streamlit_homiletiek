"""Export van één kerkdienstanalyse naar een Word-document.

Pure module zonder Streamlit-afhankelijkheden: alle UI-feedback (progressbar
e.d.) loopt via een optionele callback die de aanroeper meegeeft. Zo blijft
de documentgeneratie los te testen en te hergebruiken buiten een Streamlit-
context.

Gebruik:

    from src.utils.word_export import bouw_kerkdienstanalyse_docx
    bytes_out, bestandsnaam = bouw_kerkdienstanalyse_docx(
        analysis_id, api_handler,
        voortgang_callback=lambda f, t: print(f"{f:.0%} {t}"),
    )
"""
import io
import re
from datetime import datetime
from typing import Any, Callable
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.api.handler import APIHandler
from src.utils.analyse_tabs import TAB_NAAR_ANALYSES, TAB_VOLGORDE
# Hergebruik de centrale snake_case→leesbaar-helper; de lokale variant
# stond hier eerder als _humanize_key + _WORD_SEP_RE en is verplaatst
# naar utils.humanize_key zodat ook de Streamlit-renderers hem kunnen
# gebruiken (Noordmans-preekschets toont 'type'-labels die voorheen als
# bv. '(subversieve_wending)' rauw in de UI verschenen).
from src.utils.utils import (
    format_verse_range,
    groepeer_samengevoegde_verzen,
    humanize_key as _humanize_key,
)

# Signature voor de optionele voortgangs-callback. De aanroeper ontvangt een
# fractie (0.0 - 1.0) en een korte Nederlandse statusregel.
VoortgangCallback = Callable[[float, str], None]

# Backend-timestamps komen in UTC binnen; we tonen ze in Nederlandse tijd
# (DST-aware via zoneinfo), identiek aan dashboard.py.
_DISPLAY_TZ = ZoneInfo("Europe/Amsterdam")

# Sleutels die we bij een list-of-dicts gebruiken als subheading. Volgorde =
# prioriteit; de eerst-gematchte aanwezige sleutel wint.
_SUBHEADING_KEYS: tuple[str, ...] = (
    "titel",
    "title",
    "naam",
    "name",
    "reference",
    "heading",
    "label",
)

# Maximaal toegestaan heading-niveau in python-docx is theoretisch 9, maar
# alles boven 5 oogt scheef in Word; we cappen daarom op 5.
_MAX_HEADING_LEVEL = 5

# Drempel voor inline-vs-block rendering van string-waarden in een dict. Onder
# deze lengte (en zonder newlines) tonen we 'Sleutel: korte waarde' inline;
# daarboven krijgt de sleutel een subheading en wordt de tekst als losse
# paragraphs gerenderd (zodat meerregelige analyse-output leesbaar blijft).
_INLINE_STRING_MAX = 120

# Scheidingsteken voor paragraph-splitting binnen één string-waarde: elke
# reeks van ≥ 1 blank line telt als nieuwe alinea.
_BLANK_LINE_RE = re.compile(r"\n\s*\n")

# Tavily is de search-tool die de agent intern gebruikt; verwijzingen ernaar
# (bv. "datum onbekend (Tavily)", "via Tavily Search Tool", URL's naar
# tavily.com) horen niet in een document dat de prediker meeneemt naar
# de kerkenraad. We scrubben string-leafs op drie niveaus: parenthetische
# mentions wegknippen, losse Tavily-tokens met optionele begeleidingswoorden
# vervangen door leeg, en URL's naar tavily.com strippen. Aanvullend slaan
# we dict-keys met 'tavily' in de naam volledig over.
_TAVILY_KEY_FRAGMENT = "tavily"
_TAVILY_PAREN_RE = re.compile(
    r"\s*\([^)]*\btavily\b[^)]*\)", flags=re.IGNORECASE
)
_TAVILY_URL_RE = re.compile(r"https?://\S*tavily\S*", flags=re.IGNORECASE)
_TAVILY_PHRASE_RE = re.compile(
    r"(?:\b(?:via|uit|door|volgens|met behulp van|met)\s+)?"
    r"\bTavily(?:[\s-](?:Search(?:\s+Tool)?|zoekopdracht|zoekresultaten?|"
    r"onderzoek|tool|tools|bron|bronnen|results?|api))?\b",
    flags=re.IGNORECASE,
)
_TAVILY_DOMAIN_RE = re.compile(
    r"\btavily\.[a-z]+(?:\.[a-z]+)*(?:/\S*)?", flags=re.IGNORECASE
)
_WHITESPACE_COLLAPSE_RE = re.compile(r"[ \t]{2,}")
_DANGLING_PUNCT_RE = re.compile(
    r"\s*[—–\-,;:]\s*$|^\s*[—–\-,;:]\s*"
)

# Velden die nooit in user-facing output horen — interne diagnostische of
# database-id-kolommen die de backend in de structured_output meeneemt.
# De Streamlit-UI verbergt ze al; de generieke walker moet hetzelfde doen
# zodat ze ook niet via de 'Overig'-tab in een Word-export lekken.
_INTERNE_KEYS: frozenset[str] = frozenset({
    "source_ref_id",
})

# Sentinel-tekstwaarden die de backend invult wanneer een optioneel
# liturgisch onderdeel ontbreekt (bv. een viering met alleen een eerste
# lezing, dus geen Psalm/Epistel/Evangelie). De UI rendert zulke velden
# niet; in het Word-document zorgen ze nu nog voor 'Boek: n.v.t.' /
# 'Hoofdstuk: 0' / 'Geen evangelielezing opgegeven.'-ruis. Zie het
# voorbeelddocument kerkdienstanalyse_bethelkerk_2026-05-31.docx.
_SENTINEL_STRINGS: frozenset[str] = frozenset({
    "n.v.t.",
    "n.v.t",
    "nvt",
    "n/a",
    "n.t.b.",
    "n.t.b",
})

# Patroon voor "Geen <iets> opgegeven(.)" — komt in vele varianten voor
# in de bijbelteksten- en lezingen-output ("Geen psalmverzen opgegeven.",
# "Geen epistellezing opgegeven.", "Geen evangelielezing opgegeven."). Eén
# regex vangt alle huidige en toekomstige verschijningen, in plaats van
# elke variant met de hand op te sommen.
_GEEN_OPGEGEVEN_RE = re.compile(
    r"^\s*geen\b.*\bopgegeven\.?\s*$",
    flags=re.IGNORECASE,
)

# Patroon voor "n.v.t."-varianten met een optionele parenthetische
# toelichting erachter, bv. "n.v.t. (eigen lezingen scenario)". De
# backend gebruikt zo'n suffix om uit te leggen waarom een lezing leeg
# is; semantisch blijft het 'geen waarde' en hoort het niet in het
# document. Strenger dan een 'starts-with' match: we accepteren alleen
# n.v.t. zelf, eventueel gevolgd door één parenthetisch blok of
# afsluitende punt.
_NVT_RE = re.compile(
    r"^\s*n\.?\s*v\.?\s*t\s*\.?\s*(?:\([^)]*\))?\s*\.?\s*$",
    flags=re.IGNORECASE,
)

# Sleutels waar een numerieke `0` als sentinel ('niet opgegeven') wordt
# gebruikt. Buiten deze whitelist behouden we `0` als reguliere waarde,
# want elders kan het een echte uitkomst zijn (bv. tellingen).
_NUL_SENTINEL_KEYS: frozenset[str] = frozenset({
    "chapter",
    "hoofdstuk",
    "nummer",
})


def _is_sentinel(value: Any, key: str | None = None) -> bool:
    """True wanneer `value` een sentinel-marker is voor 'geen data'.

    De backend levert voor onbekende of overgeslagen velden vaste
    placeholder-waarden ("n.v.t.", "Geen X opgegeven.", `0` voor
    hoofdstuk-achtige sleutels). De UI verbergt ze, het Word-document
    moet hetzelfde doen.

    `key` is optioneel; alleen voor numerieke sentinels (waar `0` per
    sleutel wel/niet betekenisvol is) wordt hij geraadpleegd.
    """
    if value is None:
        return False  # None wordt al apart afgevangen door de walker
    if isinstance(value, str):
        gestript = value.strip()
        if not gestript:
            return False  # lege strings worden al elders afgevangen
        if gestript.lower() in _SENTINEL_STRINGS:
            return True
        if _NVT_RE.match(gestript):
            return True
        if _GEEN_OPGEGEVEN_RE.match(gestript):
            return True
        return False
    if isinstance(value, bool):
        return False
    if isinstance(value, int):
        if value == 0 and key in _NUL_SENTINEL_KEYS:
            return True
    return False


def _is_renderbaar(value: Any, key: str | None = None) -> bool:
    """True als `value` na sentinel-/leeg-filtering nog inhoud overhoudt.

    Wordt door `_dict_is_volledig_leeg` gebruikt om te bepalen of een
    sub-sectie helemaal weggelaten kan worden, en door `_walk_dict` om
    te beslissen of een sleutel überhaupt gerenderd hoeft te worden.
    """
    if value is None or value == "":
        return False
    if isinstance(value, str):
        # Sentinels zijn alleen "geen data" als ze óók geen zinvolle
        # andere woorden bevatten; we checken op exact-sentinel via
        # _is_sentinel zodat een verhalende tekst die toevallig 'n.v.t.'
        # bevat niet alsnog wordt weggegooid.
        return not _is_sentinel(value, key)
    if isinstance(value, dict):
        return not _dict_is_volledig_leeg(value)
    if isinstance(value, list):
        return any(_is_renderbaar(item) for item in value)
    if _is_sentinel(value, key):
        return False
    return True


def _dict_is_volledig_leeg(d: dict) -> bool:
    """True als geen enkele leaf in `d` na sentinel-filtering renderbaar is.

    Recursief: een sub-dict telt alleen als 'gevuld' wanneer hij zelf
    minstens één renderbare leaf bevat. Zo wordt een Psalm-blok met
    alleen sentinel-velden samen met zijn heading overgeslagen.
    """
    for k, v in d.items():
        if isinstance(k, str) and (
            _TAVILY_KEY_FRAGMENT in k.lower() or k in _INTERNE_KEYS
        ):
            continue
        if _is_renderbaar(v, k):
            return False
    return True


def _scrub_tavily(text: str) -> str:
    """Verwijdert Tavily-verwijzingen uit een string-leaf.

    De backend laat in fallback-meldingen soms tooling-namen lekken (bijv.
    'datum onbekend (Tavily)' of 'via Tavily Search Tool'). Voor de
    eindgebruiker zijn dat interne ruis-verwijzingen die niet in het
    geëxporteerde document horen. We strippen ze op de minst destructieve
    manier: eerst parentheticals waar 'tavily' in voorkomt verwijderen,
    daarna URL's en losse Tavily-tokens (incl. veelvoorkomende compounds
    en optionele begeleidingswoorden 'via/uit/...'), en tot slot dubbele
    whitespace en losse punctuatie aan rand of staart opruimen.
    """
    if not text:
        return text
    schoon = _TAVILY_PAREN_RE.sub("", text)
    schoon = _TAVILY_URL_RE.sub("", schoon)
    schoon = _TAVILY_DOMAIN_RE.sub("", schoon)
    schoon = _TAVILY_PHRASE_RE.sub("", schoon)
    schoon = _WHITESPACE_COLLAPSE_RE.sub(" ", schoon)
    # Losse punctuatie aan begin/einde die door het scrubben is overgebleven.
    schoon = _DANGLING_PUNCT_RE.sub("", schoon).strip()
    return schoon

def bouw_kerkdienstanalyse_docx(
    analysis_id: int,
    api_handler: APIHandler,
    voortgang_callback: VoortgangCallback | None = None,
) -> tuple[bytes, str]:
    """Bouwt een .docx met alle voltooide sub-analyses van één kerkdienstanalyse.

    Args:
        analysis_id: ID van de kerkdienstanalyse (SermonAnalysis in de backend).
        api_handler: Centrale APIHandler-instantie uit session_state.
        voortgang_callback: Optionele callback die bij elke stap wordt
            aangeroepen met (fractie 0.0-1.0, status-tekst in NL). Zo kan de
            UI een progressbar laten meelopen zonder dat deze module
            Streamlit hoeft te importeren.

    Returns:
        (bytes, bestandsnaam) — de ruwe .docx-bytes en een voorgestelde
        bestandsnaam gebaseerd op gemeente en zondagdatum.
    """
    # No-op callback zodat we verderop niet steeds op None hoeven te checken.
    cb: VoortgangCallback = voortgang_callback or (lambda f, t: None)

    cb(0.00, "Metadata ophalen...")
    sermon = api_handler.get(f"api/sermon-analyses/{analysis_id}/")

    cb(0.05, "Sub-analyses ophalen...")
    sub_raw = api_handler.get(
        "api/analysis-results",
        params={"sermon_analysis_id": analysis_id},
    )
    sub_analyses = _lijst_normaliseren(sub_raw)
    # Behoud per analyse-type alleen de meest recente run, identiek aan
    # overview.py:399-403. Soft-deleted records zijn al door de backend
    # gefilterd (analysis_result/views.py:73-76), maar bij herhaaldelijk
    # klikken op 'Opnieuw' staan er meerdere actieve records van hetzelfde
    # type in de respons. Zonder deze deduplicatie zouden bv. de
    # bijbelteksten-runs allemaal achter elkaar in het Word-document
    # verschijnen, terwijl de Streamlit-UI alleen de nieuwste toont.
    sub_analyses = _alleen_nieuwste_per_type(sub_analyses)
    # Filter analyses die in de DB als "(intern)" gemarkeerd zijn — bv.
    # `brueggemann_methode_selector` met front_end_name "Brueggemann
    # methode-selector (intern)". Dat zijn diagnostische tussenstappen
    # van de agent-pijplijn (welke methode kiezen we, welke variant
    # gebruiken we), geen output die de prediker wil zien. Schaalbaar:
    # elke toekomstige analyse met "(intern)" in de display-naam wordt
    # automatisch uitgesloten zonder extra mapping bij te werken.
    sub_analyses = _filter_interne_analyses(sub_analyses)

    cb(0.15, "Document opbouwen — kop...")
    doc = Document()
    # Footer met "Pagina X van Y", gecentreerd. Moet vóór de content, want
    # de footer geldt voor alle pagina's van de huidige sectie.
    _voeg_paginanummers_toe(doc)
    _render_kop(doc, sermon)

    # Plan eerst de groepering + bookmark-namen, zodat de inhoudsopgave
    # vooraf gerenderd kan worden met dezelfde bookmark-namen die later
    # aan de tab- en analyse-headings worden gehangen. Door de TOC zelf
    # op te bouwen (in plaats van een Word TOC-veld) hoeft Word bij het
    # openen geen velden meer bij te werken — en verschijnt de prompt
    # "Dit document bevat velden die verwijzen naar andere bestanden"
    # niet meer.
    gegroepeerd = _groepeer_per_tab(sub_analyses)
    tabs_in_volgorde: list[str] = list(TAB_VOLGORDE) + ["Overig"]
    toc_entries, tab_bookmarks, sub_bookmarks = _plan_inhoudsopgave(
        gegroepeerd, tabs_in_volgorde
    )

    # Klikbare inhoudsopgave (tab-koppen + analyse-namen = niveau 1-2).
    # Komt direct onder de kop zodat de lezer meteen kan navigeren; daarna
    # een page break zodat de eerste inhoudelijke tab op een nieuwe pagina
    # begint. Bij een lege analyse-set slaan we de TOC over — een lege
    # inhoudsopgave heeft geen meerwaarde en de fallback-melding verderop
    # vangt dat geval al af.
    if toc_entries:
        _render_inhoudsopgave(doc, toc_entries)
        doc.add_page_break()

    # Voortgang tijdens de per-analyse loop: lineair verdeeld over het
    # interval [0.20, 0.95]. Bij 0 analyses slaan we het interval simpelweg
    # over (max(...,1) voorkomt deling door 0).
    n_totaal = sum(len(lijst) for lijst in gegroepeerd.values())
    fractie_base = 0.20
    fractie_stap = (0.95 - fractie_base) / max(n_totaal, 1)

    # Volgorde van tabs is hierboven al bepaald (`tabs_in_volgorde`) zodat
    # de TOC-planning dezelfde volgorde gebruikt als de feitelijke render.
    i_globaal = 0
    # Een nieuwe analyse start in principe op een nieuwe pagina (handig bij
    # documenten met honderden pagina's). Uitzondering: de eerste analyse
    # binnen een tab blijft op dezelfde pagina als de tab-heading, anders
    # krijg je halve lege pagina's met alleen "Basis", "Verdieping" enz.
    # Elke nieuwe tab krijgt wél een eigen startpagina.
    eerste_tab = True
    for tab in tabs_in_volgorde:
        lijst = gegroepeerd.get(tab, [])
        if not lijst:
            continue
        if not eerste_tab:
            doc.add_page_break()
        eerste_tab = False
        # Bookmark op de tab-heading, zodat de overeenkomstige TOC-regel
        # er klikbaar naartoe kan springen.
        tab_heading = doc.add_heading(tab, level=1)
        bm = tab_bookmarks.get(tab)
        if bm is not None:
            _voeg_bookmark_toe(tab_heading, bm[0], bm[1])
        for idx_in_tab, sub in enumerate(lijst):
            if idx_in_tab > 0:
                # Tweede en volgende analyse binnen een tab: nieuwe pagina.
                doc.add_page_break()
            front_name = _front_end_name(sub)
            cb(
                fractie_base + i_globaal * fractie_stap,
                f"Bezig met {front_name} ({i_globaal + 1}/{n_totaal})...",
            )
            # Bookmark voor deze analyse koppelen aan de heading binnen
            # _render_analyse via de meegegeven bookmark-tuple.
            sub_bm = sub_bookmarks.get(id(sub))
            _render_analyse(doc, sub, bookmark=sub_bm)
            i_globaal += 1

    if n_totaal == 0:
        # Edge case: kerkdienstanalyse zonder enige sub-analyse. We leveren
        # nog steeds een document met de kop + expliciete melding, zodat de
        # gebruiker niet voor een onverklaarbaar lege file staat.
        doc.add_paragraph(
            "Er zijn nog geen sub-analyses om te exporteren voor deze "
            "kerkdienstanalyse."
        )

    cb(0.98, "Document opslaan...")
    buffer = io.BytesIO()
    doc.save(buffer)
    bytes_out = buffer.getvalue()

    cb(1.00, "Klaar!")
    return bytes_out, _bestandsnaam(sermon)


# ---------------------------------------------------------------------------
# Datafetch-helpers
# ---------------------------------------------------------------------------


def _lijst_normaliseren(respons: Any) -> list[dict]:
    """Normaliseert de API-respons naar een gewone list[dict].

    Django REST Framework kan bij paginatie een wrapper-dict leveren met
    sleutel 'results'; in andere gevallen komt er direct een list terug.
    Deze helper maakt er altijd een list van zodat de rest van de module
    uniform kan werken.
    """
    if isinstance(respons, list):
        return respons
    if isinstance(respons, dict) and isinstance(respons.get("results"), list):
        return respons["results"]
    return []


def _filter_interne_analyses(sub_analyses: list[dict]) -> list[dict]:
    """Verwijdert sub-analyses die in de DB als "(intern)" gemarkeerd zijn.

    De convention in `analysis_result.AnalysisType.front_end_name` is dat
    diagnostische / selector-types (zoals `brueggemann_methode_selector`)
    de suffix "(intern)" krijgen. Die rapporteren over keuzes binnen de
    agent-pijplijn — niet over de inhoud van de preek — en horen dus
    niet in het document dat de prediker meeneemt.

    Detectie via de display-naam in plaats van een hard-coded skip-lijst,
    zodat een toekomstige interne analyse automatisch wordt uitgesloten
    zonder dat we deze module hoeven bij te werken.
    """
    schoon: list[dict] = []
    for sub in sub_analyses:
        if not isinstance(sub, dict):
            schoon.append(sub)
            continue
        front = ((sub.get("analysis_type") or {}).get("front_end_name") or "")
        if "(intern)" in front.lower():
            continue
        schoon.append(sub)
    return schoon


def _alleen_nieuwste_per_type(sub_analyses: list[dict]) -> list[dict]:
    """Houd per analysis_type alleen de AnalysisResult met de hoogste id.

    Deze deduplicatie spiegelt de logica in
    ``page_navigation/analysis_results/overview.py:399-403`` zodat de
    Streamlit-UI en het Word-document dezelfde versie van elke analyse
    tonen. Records zonder een leesbare ``id`` of ``analysis_type.name``
    laten we ongemoeid in de output staan, want filteren zou data
    verliezen waar we de oorzaak nog niet van kennen.
    """
    nieuwste_per_naam: dict[str, dict] = {}
    overig: list[dict] = []
    for sub in sub_analyses:
        if not isinstance(sub, dict):
            continue
        naam = (sub.get("analysis_type") or {}).get("name")
        sub_id = sub.get("id")
        if not isinstance(naam, str) or not isinstance(sub_id, int):
            overig.append(sub)
            continue
        bestaande = nieuwste_per_naam.get(naam)
        if bestaande is None or sub_id > bestaande.get("id", -1):
            nieuwste_per_naam[naam] = sub
    return list(nieuwste_per_naam.values()) + overig


def _groepeer_per_tab(sub_analyses: list[dict]) -> dict[str, list[dict]]:
    """Verdeelt sub-analyses over de zes tabs + een 'Overig'-bak.

    Onbekende analyse-namen belanden in 'Overig' zodat we geen data
    verliezen wanneer er later in de backend een type bijkomt dat hier nog
    niet in de mapping zit. Binnen elke tab sorteren we op het `order`-veld
    van het analyse-type; bij afwezigheid vallen we terug op 99 zodat de
    onbekende types achteraan in de tab komen.
    """
    result: dict[str, list[dict]] = {tab: [] for tab in TAB_VOLGORDE}
    result["Overig"] = []

    for sub in sub_analyses:
        name = sub.get("analysis_type", {}).get("name", "")
        gevonden = False
        for tab in TAB_VOLGORDE:
            if name in TAB_NAAR_ANALYSES[tab]:
                result[tab].append(sub)
                gevonden = True
                break
        if not gevonden:
            result["Overig"].append(sub)

    for lijst in result.values():
        lijst.sort(key=lambda s: s.get("analysis_type", {}).get("order", 99))

    return result


def _front_end_name(sub: dict) -> str:
    """Geeft de display-naam voor een sub-analyse, gescrubd van Tavily-verwijzingen.

    Zes analyse-types staan in de DB met `front_end_name` in de vorm
    "Gemeente-spiritualiteit (Tavily)" — handig in de UI om aan te geven
    dat de pijplijn een externe zoektool gebruikt, maar verkeerd in een
    document dat de prediker meeneemt naar de kerkenraad. We strippen de
    parenthetische Tavily-mention via dezelfde scrub-pipeline die ook
    voor leaf-tekst geldt; zo blijven zowel de inhoudsopgave als de
    sectiekoppen schoon.
    """
    at = sub.get("analysis_type") or {}
    naam = (
        at.get("front_end_name")
        or _humanize_key(at.get("name", ""))
        or "Onbekende analyse"
    )
    schoon = _scrub_tavily(naam).strip() if naam else naam
    return schoon or "Onbekende analyse"


# ---------------------------------------------------------------------------
# Rendering: kop + per analyse
# ---------------------------------------------------------------------------


def _render_kop(doc, sermon: dict) -> None:
    """Rendert titel + metadata bovenaan het document."""
    doc.add_heading(_kop_titel(sermon), level=0)

    gemeente = (sermon.get("church") or {}).get("name", "")
    sermon_date_iso = sermon.get("sermon_date", "")
    aanmaak = _format_aanmaak_label(sermon.get("created_at"))

    if gemeente:
        _metadata_regel(doc, "Gemeente", gemeente)
    if sermon_date_iso:
        _metadata_regel(doc, "Zondagdatum", _format_sermon_date(sermon_date_iso))
    if aanmaak:
        _metadata_regel(doc, "Aangemaakt", aanmaak)

    # Witregel tussen kop en eerste tabblad.
    doc.add_paragraph()


def _metadata_regel(doc, label: str, waarde: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(waarde)


def _kop_titel(sermon: dict) -> str:
    """Kopt op basis van de title van de analyse; valt terug op gemeente+datum."""
    titel = (sermon.get("title") or "").strip()
    if titel:
        return titel
    gemeente = (sermon.get("church") or {}).get("name", "")
    sermon_date = _format_sermon_date(sermon.get("sermon_date", ""))
    return f"{gemeente} — {sermon_date}".strip(" —")


def _format_sermon_date(iso_datum: str) -> str:
    # Zondagdatum komt in 'YYYY-MM-DD'-formaat van de API.
    try:
        return datetime.strptime(iso_datum, "%Y-%m-%d").strftime("%d-%m-%Y")
    except ValueError:
        return iso_datum


def _format_aanmaak_label(created_at: str | None) -> str | None:
    """Converteert ISO-8601 UTC-timestamp naar 'dd-mm-yyyy HH:MM' in NL-tijd.

    Retourneert None wanneer het veld ontbreekt of niet geparsed kan worden;
    de aanroeper laat de regel in dat geval weg.
    """
    if not created_at:
        return None
    try:
        dt_utc = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
        return dt_utc.astimezone(_DISPLAY_TZ).strftime("%d-%m-%Y %H:%M")
    except ValueError:
        return None


def _render_analyse(
    doc,
    sub: dict,
    bookmark: tuple[int, str] | None = None,
) -> None:
    """Rendert één sub-analyse: heading + body (walk) of een lege-melding.

    We filteren niet op een `status`-veld: de bestaande renderers in
    page_navigation/analysis_results/analyses/ doen dat ook niet, en in de
    praktijk blijkt het statusveld vaak niet op de verwachte sentinel-
    waarde ('completed') te staan terwijl `result` wél bruikbare data
    bevat. Presence van `result` is de enige betrouwbare indicator.

    `bookmark` is een optionele (id, naam)-tuple die op de analyse-heading
    wordt gehangen zodat de inhoudsopgave er klikbaar naartoe kan
    springen.
    """
    heading = doc.add_heading(_front_end_name(sub), level=2)
    if bookmark is not None:
        _voeg_bookmark_toe(heading, bookmark[0], bookmark[1])

    result = sub.get("result")
    if _is_leeg_resultaat(result):
        p = doc.add_paragraph()
        p.add_run("— nog geen resultaat beschikbaar.").italic = True
        return

    # Per-analyse dispatch: bijbelteksten heeft een eigen renderer omdat
    # de generieke walker de geneste verzen-structuur niet zinvol kan
    # weergeven (dubbele nummering + Engelse interne sleutels). Andere
    # analyses blijven via de generieke walker lopen.
    analysis_name = (sub.get("analysis_type") or {}).get("name", "")
    if analysis_name == "bijbelteksten":
        _render_bijbelteksten(doc, result, heading_level=3)
        return

    _walk(result, doc, heading_level=3)


def _is_leeg_resultaat(result: Any) -> bool:
    """True wanneer er niks zinnigs in `result` staat om te renderen."""
    if result is None:
        return True
    if isinstance(result, (str, list, dict)) and len(result) == 0:
        return True
    return False


# ---------------------------------------------------------------------------
# Generieke JSON → docx walker
# ---------------------------------------------------------------------------


def _walk(
    value: Any,
    doc,
    heading_level: int = 3,
    skip_keys: set[str] | None = None,
) -> None:
    """Recursieve JSON-walker.

    - dict  → sleutels als "Key: value"-paragraph of als subheading+recursie.
    - list  → list-of-dict → secties, list-of-primitives → bullets.
    - str   → paragraph (met dubbele-newline-splitting voor lange teksten).
    - bool  → "Ja"/"Nee" paragraph.
    - None of ""  → overgeslagen.
    - overig (getal) → str(x) paragraph.
    """
    if value is None or value == "":
        return
    if isinstance(value, dict):
        _walk_dict(value, doc, heading_level, skip_keys or set())
    elif isinstance(value, list):
        _walk_list(value, doc, heading_level)
    elif isinstance(value, bool):
        doc.add_paragraph("Ja" if value else "Nee")
    elif isinstance(value, str):
        _add_paragraph_multiline(doc, value)
    else:
        # Niet-string primitieven (getallen) bevatten geen Tavily-tekst, dus
        # gewoon doorgeven aan str().
        doc.add_paragraph(str(value))


def _walk_dict(
    value: dict,
    doc,
    heading_level: int,
    skip_keys: set[str],
) -> None:
    """Doorloopt een dict in insertion-order.

    Primitieve waarden (str/number/bool) worden als "Key: value"-paragraph
    gerenderd; nested dicts/lists krijgen een subheading en recursieve
    verwerking. Sleutels in `skip_keys` worden overgeslagen — handig wanneer
    de aanroepende _walk_list de titel al als heading heeft gerenderd.
    """
    for key, val in value.items():
        if key in skip_keys:
            continue
        # Sla velden over waarvan de sleutel zelf naar Tavily verwijst
        # (bv. 'tavily_query', 'tavily_results'). Die zijn diagnostisch
        # voor de backend en horen niet in het Word-document.
        if isinstance(key, str) and _TAVILY_KEY_FRAGMENT in key.lower():
            continue
        # Interne database-id's en andere diagnostische sleutels (bv.
        # source_ref_id in bijbelteksten) horen evenmin in user-facing
        # output. De dedicated bijbelteksten-renderer bedient die data
        # nu zelf; deze skip is een vangnet voor edge-cases waarin
        # diezelfde sleutels alsnog via _walk binnenkomen.
        if isinstance(key, str) and key in _INTERNE_KEYS:
            continue
        label = _humanize_key(key)
        if isinstance(val, (dict, list)):
            # Sla zowel heading als recursieve walk over wanneer de
            # geneste structuur na sentinel-filtering niets renderbaars
            # overhoudt — anders blijven kale headings achter (bv. een
            # Psalm-kop zonder enige inhoud).
            if not _is_renderbaar(val, key):
                continue
            niveau = min(heading_level, _MAX_HEADING_LEVEL)
            doc.add_heading(label, level=niveau)
            _walk(val, doc, heading_level + 1)
        elif val is None or val == "":
            # Lege velden overslaan voorkomt eindeloze "Key: "-regels zonder
            # inhoud in documenten waar veel optionele velden leegblijven.
            continue
        elif _is_sentinel(val, key):
            # Sentinel-leaves ("n.v.t.", "Geen X opgegeven.", `0` voor
            # hoofdstuk-achtige sleutels) markeren afwezigheid; de UI
            # verbergt ze, dus wij ook.
            continue
        elif isinstance(val, str) and _is_block_tekst(val):
            # Lange of meerregelige strings krijgen een eigen subheading +
            # losse paragraphs, anders verdwijnen de alinea-grenzen uit de
            # model-output in één lange inline-regel.
            niveau = min(heading_level, _MAX_HEADING_LEVEL)
            schoon = _scrub_tavily(val)
            if not schoon:
                continue
            doc.add_heading(label, level=niveau)
            _add_paragraph_multiline(doc, schoon)
        else:
            # Inline 'Sleutel: waarde' — string-waardes scrubben; als er na
            # het scrubben niets overblijft slaan we de hele regel over zodat
            # we geen 'Datum: '-regels krijgen waar alleen 'Tavily' stond.
            if isinstance(val, str):
                schoon = _scrub_tavily(val)
                if not schoon:
                    continue
                val = schoon
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            _append_value_run(p, val)


def _walk_list(value: list, doc, heading_level: int) -> None:
    if not value:
        return

    # Drie gevallen:
    # 1) alle items zijn dicts            → elk item als sectie met subheading
    # 2) alle items zijn primitieven      → bullet-list
    # 3) gemengd (zeldzaam in onze data)  → per-item recursieve fallback
    if all(isinstance(item, dict) for item in value):
        # Als géén van de items een titel-achtige sleutel heeft, voegen
        # we geen valse "1.", "2."-koppen toe. Dat patroon trad nu op
        # bij Valkuilen (`{valkuil, risico, advies}` zonder titel-veld)
        # en leverde een nodeloos hiërarchisch niveau op terwijl de items
        # zelf al duidelijk gestructureerd zijn. Witregel tussen items
        # houdt de visuele scheiding zonder pseudo-kop.
        subheadings = [_subheading_van_dict(item) for item in value]
        heeft_titels = any(sub for _, sub in subheadings)
        niveau = min(heading_level, _MAX_HEADING_LEVEL)
        if not heeft_titels:
            for i, item in enumerate(value):
                if i > 0:
                    # Witregel scheidt opeenvolgende items zonder kop.
                    doc.add_paragraph()
                _walk(item, doc, heading_level)
            return
        for i, (item, (subkey, subwaarde)) in enumerate(
            zip(value, subheadings), start=1
        ):
            if subwaarde:
                doc.add_heading(subwaarde, level=niveau)
                # De gebruikte sleutel overslaan, anders wordt hij hieronder
                # nóg een keer als "Titel: ..."-paragraph gerenderd.
                _walk(item, doc, heading_level + 1, skip_keys={subkey})
            else:
                doc.add_heading(f"{i}.", level=niveau)
                _walk(item, doc, heading_level + 1)
        return

    if all(not isinstance(item, (dict, list)) for item in value):
        for item in value:
            if item is None or item == "":
                continue
            if _is_sentinel(item):
                # Sentinel-bullets zoals "n.v.t." in een lijst zijn ruis;
                # dezelfde regel als voor dict-leaves toepassen.
                continue
            if isinstance(item, bool):
                tekst = "Ja" if item else "Nee"
            else:
                tekst = _scrub_tavily(str(item))
                # Bullet-items die volledig uit een Tavily-verwijzing bestonden
                # worden hierdoor leeg en moeten worden overgeslagen — anders
                # krijgen we een lege bullet in de lijst.
                if not tekst:
                    continue
            doc.add_paragraph(tekst, style="List Bullet")
        return

    # Gemengde lijst: per item recursief verwerken zonder lijstopmaak.
    for item in value:
        _walk(item, doc, heading_level)


def _subheading_van_dict(item: dict) -> tuple[str | None, str | None]:
    """Zoekt de eerste bruikbare titel-achtige sleutel in een dict.

    Retourneert (sleutel, waarde) zodat de caller de sleutel aan
    `skip_keys` kan toevoegen. Beide None als er geen geschikte sleutel is.
    """
    for key in _SUBHEADING_KEYS:
        val = item.get(key)
        if isinstance(val, str) and val.strip():
            # Subheadings worden als heading in het document gerenderd, dus
            # ook hier eerst de Tavily-mentions strippen voordat we hem
            # promoveren — anders verschijnt 'Tavily' alsnog als kop.
            schoon = _scrub_tavily(val.strip())
            if schoon:
                return key, schoon
    return None, None


def _append_value_run(p, value: Any) -> None:
    if isinstance(value, bool):
        p.add_run("Ja" if value else "Nee")
    else:
        p.add_run(str(value))


def _add_paragraph_multiline(doc, text: str) -> None:
    """Splitst een string op blank lines en maakt er aparte paragraphs van.

    Zo blijven gestructureerde analyse-resultaten met meerdere alinea's ook
    in Word leesbaar: de model-output gebruikt vaak '\\n\\n' als alinea-
    scheider, wat python-docx anders als één grote paragraph zou behandelen.
    Per stuk wordt nog een keer door _scrub_tavily gehaald, zodat ook
    callers die deze helper rechtstreeks aanroepen geen Tavily-mentions
    door kunnen lekken.
    """
    text = text.strip()
    if not text:
        return
    stukken = _BLANK_LINE_RE.split(text)
    for stuk in stukken:
        gestript = _scrub_tavily(stuk.strip())
        if not gestript:
            continue
        doc.add_paragraph(gestript)


def _is_block_tekst(tekst: str) -> bool:
    """True als een string beter als block-paragraph gerenderd kan worden.

    Criterium: bevat een newline (dus meerdere regels) óf is langer dan
    _INLINE_STRING_MAX tekens (zonder newlines). Kort+eenregelig blijft
    inline in 'Sleutel: waarde'-vorm.
    """
    if "\n" in tekst:
        return True
    return len(tekst) > _INLINE_STRING_MAX


# ---------------------------------------------------------------------------
# Bijbelteksten — dedicated renderer
# ---------------------------------------------------------------------------
#
# De generieke walker gaf voor elke bijbeltekst-passage:
#   1. een nodeloze "Verses"-kop (Engels, uit `verses`-sleutel),
#   2. per vers een "1.", "2."-sectie + alle keys als "Number: 1",
#      "Modern text: ...", "Source ref id: 7620" inline,
#   3. de Hebreeuwse brontekst per vers apart (zelfs bij identieke tekst
#      in opeenvolgende verzen),
#   4. ten slotte een hangende "Chapter: 17"-regel onderaan elke passage.
#
# De UI in `page_navigation/analysis_results/analyses/bijbelteksten.py`
# laat al die ruis weg; de Word-export moet hetzelfde detailniveau
# leveren. Hieronder een compacte mirror van de UI-pijplijn:
#   _spreid_source_text → groepeer_samengevoegde_verzen
#                       → _cluster_op_source_ref → render
#
# Helpers zijn bewust lokaal gehouden ipv import uit de UI-module: dat
# zou een UI→utils reverse-dependency zijn. Een DRY-extractie naar
# `src/utils/utils.py` kan in een vervolg-feature.


def _is_hebrew(text: str) -> bool:
    return any("֐" <= ch <= "׿" for ch in text)


def _spreid_source_text(verses: list[dict]) -> list[dict]:
    """Vul lege source_text-verzen met de brontekst van een buur met
    identieke modern_text. Spiegel van de UI-helper met dezelfde naam.

    Zonder deze stap zou `groepeer_samengevoegde_verzen` v13 (gevuld) en
    v14 (leeg) als 'verschillend' beschouwen waardoor de prediker
    dezelfde Dutch-tekst twee keer onder elkaar zou krijgen.
    """
    schoon = [dict(v) if isinstance(v, dict) else v for v in verses]

    def _bron(v: dict) -> str:
        return (v.get("source_text") or "").strip()

    def _modern(v: dict) -> str:
        return (v.get("modern_text") or "").rstrip()

    for i in range(1, len(schoon)):
        cur, prev = schoon[i], schoon[i - 1]
        if (
            isinstance(cur, dict) and isinstance(prev, dict)
            and _modern(cur) and _modern(cur) == _modern(prev)
            and not _bron(cur) and _bron(prev)
        ):
            cur["source_text"] = prev["source_text"]
    for i in range(len(schoon) - 2, -1, -1):
        cur, nxt = schoon[i], schoon[i + 1]
        if (
            isinstance(cur, dict) and isinstance(nxt, dict)
            and _modern(cur) and _modern(cur) == _modern(nxt)
            and not _bron(cur) and _bron(nxt)
        ):
            cur["source_text"] = nxt["source_text"]
    return schoon


def _verzamel_bron_paren(
    nummers: list[Any], verses: list[dict]
) -> list[tuple[str, Any]]:
    """Verzamel unieke (source_text, source_ref_id)-paren in voorkomvolgorde.

    Mirror van de UI-helper: één Hebreeuws/Grieks blok per unieke bron-
    regel binnen een groep gedeelde Dutch-tekst.
    """
    paren: list[tuple[str, Any]] = []
    gezien: set[tuple[str, Any]] = set()
    for n in nummers:
        v = next(
            (v for v in verses if isinstance(v, dict) and v.get("number") == n),
            None,
        )
        if v is None:
            continue
        src = (v.get("source_text") or "").strip()
        if not src:
            continue
        sleutel = (src, v.get("source_ref_id"))
        if sleutel in gezien:
            continue
        gezien.add(sleutel)
        paren.append(sleutel)
    return paren


def _cluster_op_source_ref(
    groepen: list[dict], verses: list[dict]
) -> list[dict]:
    """Voeg opeenvolgende groepen die naar dezelfde BHS-rij verwijzen
    samen tot één super-cluster. Mirror van de UI-helper met dezelfde naam.

    Voorkomt dat een gedeelde Hebreeuwse regel (bv. opschrift+vers in
    Psalmen) twee keer onder elkaar verschijnt.
    """
    ref_per_nummer: dict[Any, Any] = {}
    for v in verses:
        if not isinstance(v, dict):
            continue
        ref_per_nummer[v.get("number")] = v.get("source_ref_id")

    super_clusters: list[dict] = []
    for groep in groepen:
        nummers = groep.get("numbers") or []
        ref_ids = {ref_per_nummer.get(n) for n in nummers}
        gedeelde_ref = ref_ids.pop() if len(ref_ids) == 1 else None
        bron_paren = _verzamel_bron_paren(nummers, verses)

        if (
            super_clusters
            and gedeelde_ref is not None
            and len(super_clusters[-1]["bron_paren"]) == 1
            and super_clusters[-1]["bron_paren"][0][1] == gedeelde_ref
            and super_clusters[-1]["bron_paren"][0][1] is not None
        ):
            super_clusters[-1]["sub_groepen"].append(groep)
            continue

        super_clusters.append(
            {"sub_groepen": [groep], "bron_paren": bron_paren}
        )
    return super_clusters


def _render_bijbelteksten(doc, result: Any, heading_level: int = 3) -> None:
    """Speciale renderer voor de bijbelteksten-analyse.

    Vermijdt de drie problemen van de generieke walker:
    1. dubbele nummering (`_walk_list` "1." + dict-leaf "Number: 1"),
    2. Engelse interne sleutels ("Modern text", "Source ref id", "Chapter")
       die via `humanize_key` slechts capitaliseren zonder vertalen,
    3. brontekst die voor élk vers apart verschijnt ook bij identieke
       Hebreeuwse rijen.

    `result` is een lijst van scripture-passages — zelfde shape als de
    UI verwacht (zie `bijbelteksten.py:182`). Voor de alternatieve dict-
    shape (oude data, bv. sermon 215 met `boek`/`verzen`/`tekst_bron`)
    valt de UI terug op "Nog geen bijbeltekst beschikbaar"; wij doen
    hetzelfde zodat de uitvoer voorspelbaar is en niet plotseling een
    bulk verbose dict-walk oplevert. Wanneer alle passages na filtering
    leeg blijken, tonen we een italic placeholder zodat de "Bijbelteksten"-
    kop nooit wees blijft staan.
    """
    if not isinstance(result, list):
        # Alt-shape (dict) of legacy/error-shape: zelfde gedrag als de UI.
        # Beter een duidelijke melding dan een verwarrende fallback-walk
        # met Engelstalige sleutels en interne id's.
        p = doc.add_paragraph()
        p.add_run("— geen renderbare bijbeltekst beschikbaar.").italic = True
        return

    niveau = min(heading_level, _MAX_HEADING_LEVEL)
    iets_gerenderd = False
    for scripture in result:
        if not isinstance(scripture, dict):
            continue
        ref = (scripture.get("reference") or "").rstrip(".").strip()
        verses = scripture.get("verses") or []
        if not isinstance(verses, list):
            verses = []

        # Overslaan wanneer deze passage geen renderbare verzen bevat —
        # dan zou alleen een kale referentie-heading achterblijven.
        if not any(
            isinstance(v, dict) and (v.get("modern_text") or "").strip()
            for v in verses
        ):
            continue

        if ref:
            doc.add_heading(ref, level=niveau)

        verses = _spreid_source_text(verses)
        groepen = groepeer_samengevoegde_verzen(
            verses, text_fields=("modern_text",)
        )
        super_clusters = _cluster_op_source_ref(groepen, verses)

        for cluster in super_clusters:
            for groep in cluster["sub_groepen"]:
                label = format_verse_range(groep["numbers"])
                modern_text = (groep.get("modern_text") or "").rstrip()
                if not modern_text:
                    continue
                schoon = _scrub_tavily(modern_text)
                if not schoon:
                    continue
                p = doc.add_paragraph()
                # Bold versbereik vooraan (bv. "13-14. "), gevolgd door
                # de Dutch-tekst. Dit vervangt zowel de "1."-pseudo-kop
                # als de inline "Number: 1"-regel uit de generieke walker.
                p.add_run(f"{label}. ").bold = True
                p.add_run(schoon)
                iets_gerenderd = True

            # Eén italic-blok per unieke bron-regel in het cluster — zo
            # zien opeenvolgende Dutch-versen die naar dezelfde BHS-rij
            # mappen die rij maar één keer.
            for source_text, _ in cluster["bron_paren"]:
                schone_bron = source_text.strip()
                if not schone_bron:
                    continue
                bron_p = doc.add_paragraph()
                if _is_hebrew(schone_bron):
                    # Right-to-left tekstrichting voor Hebreeuws zodat
                    # Word de regel correct uitlijnt. python-docx heeft
                    # geen high-level API hiervoor; we zetten `w:bidi`
                    # op de paragraph-properties.
                    pPr = bron_p._element.get_or_add_pPr()
                    bidi = OxmlElement("w:bidi")
                    pPr.append(bidi)
                bron_run = bron_p.add_run(schone_bron)
                bron_run.italic = True

    if not iets_gerenderd:
        # Lege of volledig gefilterde resultaten leveren anders alleen de
        # "Bijbelteksten"-kop op uit `_render_analyse`, wat eruitziet als
        # een verdwenen sectie. Eén italic regel maakt zichtbaar dat er
        # niets renderbaars in de analyse zit (bv. een mislukte agent-run).
        p = doc.add_paragraph()
        p.add_run("— geen renderbare bijbeltekst beschikbaar.").italic = True


# ---------------------------------------------------------------------------
# Inhoudsopgave, bookmarks en paginanummers (raw OOXML)
# ---------------------------------------------------------------------------
#
# python-docx heeft geen high-level API voor veldcodes, hyperlinks naar
# bookmarks of bookmark-tags, dus we injecteren de benodigde OOXML-elementen
# rechtstreeks. Het "w:fldChar begin / instrText / separate / end"-patroon
# is de standaard-constructie die Word gebruikt voor PAGE/NUMPAGES in de
# footer.
#
# We gebruiken bewust géén Word TOC-veld: dat veld dwingt Word ertoe bij het
# openen alle velden bij te werken, wat de prompt "Dit document bevat velden
# die verwijzen naar andere bestanden. Wilt u ze bijwerken?" oplevert. In
# plaats daarvan rendert deze module de inhoudsopgave zelf als een lijst van
# klikbare hyperlinks naar interne bookmarks. Die hyperlinks zijn geen
# velden, dus Word toont geen prompt.

# Hyperlink-styling: blauw + onderstreept, rechtstreeks op de run gezet
# zodat we niet afhankelijk zijn van het bestaan van een 'Hyperlink'-style
# in de standaard python-docx Document().
_HYPERLINK_KLEUR = "0563C1"


def _veld_invoegen(run, instructie: str, placeholder: str = "") -> None:
    """Injecteert een Word-veldcode in een run.

    Word berekent de zichtbare tekst op basis van `instructie` (bv. 'PAGE'
    of 'NUMPAGES'). PAGE/NUMPAGES worden door Word lazily berekend zodra
    het document gepagineerd wordt; daar is geen `w:updateFields`-setting
    voor nodig en dus ook geen prompt bij openen.
    """
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instructie

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = placeholder

    eind = OxmlElement("w:fldChar")
    eind.set(qn("w:fldCharType"), "end")

    r_elem = run._element
    r_elem.append(begin)
    r_elem.append(instr)
    r_elem.append(separate)
    r_elem.append(placeholder_t)
    r_elem.append(eind)


def _voeg_paginanummers_toe(doc) -> None:
    """Zet 'Pagina X van Y' gecentreerd in de footer van de eerste sectie."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run("Pagina ")
    run_page = p.add_run()
    _veld_invoegen(run_page, "PAGE")
    p.add_run(" van ")
    run_total = p.add_run()
    _veld_invoegen(run_total, "NUMPAGES")


def _voeg_bookmark_toe(paragraph, bookmark_id: int, bookmark_naam: str) -> None:
    """Plaatst `<w:bookmarkStart>`/`<w:bookmarkEnd>` rond een paragraph.

    De bookmark omsluit de hele paragraph-inhoud, zodat een hyperlink met
    `w:anchor` naar deze bookmark direct naar de heading scrollt. De
    `bookmarkStart` moet ná `w:pPr` komen (anders accepteert Word het
    bestand niet meer); we vinden of voegen `w:pPr` als eerste kindelement.
    """
    p_elem = paragraph._element
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_naam)
    eind = OxmlElement("w:bookmarkEnd")
    eind.set(qn("w:id"), str(bookmark_id))

    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p_elem.insert(0, start)
    p_elem.append(eind)


def _plan_inhoudsopgave(
    gegroepeerd: dict[str, list[dict]],
    tabs_in_volgorde: list[str],
) -> tuple[
    list[tuple[int, str, int, str]],
    dict[str, tuple[int, str]],
    dict[int, tuple[int, str]],
]:
    """Plant de inhoudsopgave en wijst bookmarks toe.

    Returnt een tuple `(toc_entries, tab_bookmarks, sub_bookmarks)`:

    - `toc_entries` is de lijst die door `_render_inhoudsopgave` gerenderd
      wordt: `(level, label, bookmark_id, bookmark_naam)` per regel,
      level 1 voor tabs en level 2 voor analyses.
    - `tab_bookmarks` mapt tabnaam → `(bookmark_id, bookmark_naam)` zodat
      de tab-heading later met dezelfde bookmark gemarkeerd kan worden.
    - `sub_bookmarks` doet hetzelfde, maar per sub-analyse, gekeyd op
      `id(sub_dict)` — dat is uniek voor de levensduur van de generatie
      en vermijdt dat we ervan uit moeten gaan dat sub-analyses een
      stabiele eigen identifier hebben in dit pad.

    Bookmark-namen volgen het simpele patroon `_bm_<n>`: stabiel, kort en
    bevat geen tekens die Word lastig vindt.
    """
    toc_entries: list[tuple[int, str, int, str]] = []
    tab_bookmarks: dict[str, tuple[int, str]] = {}
    sub_bookmarks: dict[int, tuple[int, str]] = {}

    teller = 0
    for tab in tabs_in_volgorde:
        lijst = gegroepeerd.get(tab, [])
        if not lijst:
            continue
        teller += 1
        bm = (teller, f"_bm_{teller}")
        tab_bookmarks[tab] = bm
        toc_entries.append((1, tab, bm[0], bm[1]))
        for sub in lijst:
            teller += 1
            bm = (teller, f"_bm_{teller}")
            sub_bookmarks[id(sub)] = bm
            toc_entries.append((2, _front_end_name(sub), bm[0], bm[1]))

    return toc_entries, tab_bookmarks, sub_bookmarks


def _render_inhoudsopgave(
    doc, toc_entries: list[tuple[int, str, int, str]]
) -> None:
    """Rendert een vooraf opgebouwde, klikbare inhoudsopgave (geen TOC-veld).

    Elke entry wordt een paragraph met één `<w:hyperlink w:anchor="...">`
    die naar de bijbehorende bookmark op een tab- of analyse-heading
    verwijst. Doordat dit gewone hyperlinks zijn (geen veldcodes), heeft
    Word geen reden om bij het openen velden bij te werken — en dus geen
    "Dit document bevat velden..."-prompt.

    Het label 'Inhoudsopgave' is bewust géén Heading 1/2: anders zou de
    eigen heading van de inhoudsopgave als eerste item in de TOC
    verschijnen wanneer iemand later toch een TOC-veld wil gebruiken.
    """
    label = doc.add_paragraph()
    run = label.add_run("Inhoudsopgave")
    run.bold = True
    run.font.size = Pt(18)

    # Lege regel tussen label en de feitelijke TOC-inhoud.
    doc.add_paragraph()

    for level, tekst, _bm_id, bm_naam in toc_entries:
        _render_toc_regel(doc, tekst, bm_naam, level)


def _render_toc_regel(doc, tekst: str, bookmark_naam: str, level: int) -> None:
    """Rendert één klikbare regel in de inhoudsopgave.

    Niveau 2 (analyse-namen) krijgt een lichte inspring zodat de
    hiërarchie zichtbaar blijft. We gebruiken `Pt(18)` als inspring —
    dat is visueel ongeveer een tab-stop en past bij de puntgrootte van
    body text.
    """
    p = doc.add_paragraph()
    if level >= 2:
        p.paragraph_format.left_indent = Pt(18)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_naam)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), _HYPERLINK_KLEUR)
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = tekst
    r.append(t)

    hyperlink.append(r)
    p._element.append(hyperlink)


# ---------------------------------------------------------------------------
# Bestandsnaam
# ---------------------------------------------------------------------------


def _bestandsnaam(sermon: dict) -> str:
    """Levert een bestandsveilige naam o.b.v. gemeente + zondagdatum."""
    gemeente_raw = (sermon.get("church") or {}).get("name", "kerkdienst")
    sermon_date = sermon.get("sermon_date", "")
    # Vervang alles behalve letters/cijfers/underscore/koppelteken door '_',
    # lowercase, en trim trailing underscores — voorkomt problemen op
    # Windows-filesystems en ziet er netjes uit.
    veilig = re.sub(r"[^\w\-]+", "_", gemeente_raw.lower()).strip("_")
    if not veilig:
        veilig = "kerkdienst"
    if sermon_date:
        return f"kerkdienstanalyse_{veilig}_{sermon_date}.docx"
    return f"kerkdienstanalyse_{veilig}.docx"
